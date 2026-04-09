from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from datetime import datetime
from pydantic import BaseModel, Field
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document

from backend.objections import load_rules


BASE_DIR = Path(__file__).resolve().parent.parent
FRONTEND_DIR = BASE_DIR / "frontend"
RULES, RULES_SOURCE = load_rules()

app = FastAPI(title="Delhi HC Defect Checker")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory=str(FRONTEND_DIR)), name="static")


@app.get("/")
def home() -> FileResponse:
    return FileResponse(FRONTEND_DIR / "index.html")


def extract_text_from_pdf(raw_bytes: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(raw_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages).lower()
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not read PDF: {exc}") from exc


def check_objections(text: str) -> List[Dict]:
    hits = []
    for rule in RULES:
        matched_keywords = [kw for kw in rule.keywords if kw in text]
        if matched_keywords:
            score = round((len(matched_keywords) / len(rule.keywords)) * 100)
            hits.append(
                {
                    "code": rule.code,
                    "title": rule.title,
                    "confidence": score,
                    "matched_keywords": matched_keywords,
                    "guidance": rule.guidance,
                }
            )
    hits.sort(key=lambda x: x["confidence"], reverse=True)
    return hits


@app.post("/api/check")
async def check_pdf(file: UploadFile = File(...)) -> Dict:
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed.")

    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    text = extract_text_from_pdf(raw)
    detected = check_objections(text)

    return {
        "file_name": file.filename,
        "pages_scanned": max(text.count("\n"), 1),
        "total_rules_checked": len(RULES),
        "possible_defects_found": len(detected),
        "rules_source": RULES_SOURCE,
        "defects": detected,
    }


class DefectItem(BaseModel):
    code: int
    title: str
    confidence: int
    guidance: str
    date_marked: Optional[str] = None
    date_removed: Optional[str] = None
    marked_on: Optional[str] = None


class ReportPayload(BaseModel):
    diary_no: str = "N/A"
    case_title: str = "N/A"
    generated_on: str = "N/A"
    file_name: str
    defects: List[DefectItem]
    now_display: Optional[str] = None
    max_main_rows: int = Field(default=3, ge=1, le=20)


@app.post("/api/export/pdf")
def export_pdf(payload: ReportPayload):
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    now_display = payload.now_display or datetime.now().strftime("%d/%m/%Y %I:%M %p")

    y = 805
    c.setFont("Helvetica", 12)
    c.drawString(60, y, "Dear Sir/Madam")
    y -= 28
    c.setFont("Helvetica", 11)
    c.drawString(90, y, "Please check defect(s) marked against")
    y -= 18
    c.drawString(90, y, f"Diary No:{payload.diary_no} in the matter")
    y -= 18
    c.drawString(90, y, payload.case_title[:75])
    y -= 26

    # Table layout (similar to screenshot)
    left = 60
    col1 = left
    col2 = left + 45
    col3 = left + 320
    col4 = left + 430
    right = 545

    def draw_row(ypos: int, sl: str, defect: str, marked: str, removed: str):
        c.setFont("Helvetica", 8.7)
        c.drawString(col1 + 4, ypos, sl)
        c.drawString(col2 + 4, ypos, defect[:85])
        c.drawString(col3 + 4, ypos, marked)
        c.drawString(col4 + 4, ypos, removed)

    # Header
    row_h = 18
    c.setLineWidth(1)
    c.rect(left, y - (row_h * 2), right - left, row_h * 2, stroke=1, fill=0)
    c.line(col2, y, col2, y - (row_h * 2))
    c.line(col3, y, col3, y - (row_h * 2))
    c.line(col4, y, col4, y - (row_h * 2))
    c.setFont("Helvetica-Bold", 9)
    c.drawString(col1 + 4, y - 14, "SlNo.")
    c.drawString(col2 + 4, y - 14, "Defects marked during Scrutiny")
    c.drawString(col3 + 4, y - 8, "Date of")
    c.drawString(col3 + 4, y - 18, "Defects Marked")
    c.drawString(col4 + 4, y - 8, "Date of")
    c.drawString(col4 + 4, y - 18, "Defect Removed")
    y -= row_h * 2

    main = payload.defects[: payload.max_main_rows]
    other = payload.defects[payload.max_main_rows :]

    # Main rows
    for idx, d in enumerate(main, start=1):
        if y < 140:
            c.showPage()
            y = 805
        c.rect(left, y - row_h, right - left, row_h, stroke=1, fill=0)
        c.line(col2, y, col2, y - row_h)
        c.line(col3, y, col3, y - row_h)
        c.line(col4, y, col4, y - row_h)
        draw_row(
            y - 12,
            str(idx),
            f"({d.code}) - {d.title}",
            d.date_marked or payload.generated_on or "",
            d.date_removed or "",
        )
        y -= row_h

    # Any Other Defects header
    if other:
        c.rect(left, y - row_h, right - left, row_h, stroke=1, fill=0)
        c.setFont("Helvetica-Bold", 9)
        c.drawString(left + 4, y - 12, "Any Other Defects")
        c.drawString(col4 + 4, y - 12, "Marked On")
        y -= row_h
        for j, d in enumerate(other, start=len(main) + 1):
            if y < 140:
                c.showPage()
                y = 805
            c.rect(left, y - row_h, right - left, row_h, stroke=1, fill=0)
            c.line(col4, y, col4, y - row_h)
            c.setFont("Helvetica", 8.7)
            c.drawString(left + 4, y - 12, f"{j}")
            c.drawString(col2 + 4, y - 12, f"({d.code}) - {d.title}"[:92])
            c.drawString(col4 + 4, y - 12, d.marked_on or d.date_marked or payload.generated_on or "")
            y -= row_h

    y -= 14
    c.setFont("Helvetica", 11)
    c.drawString(60, y, f"Date :{now_display}")
    c.save()
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="defect_report.pdf"'},
    )


@app.post("/api/export/docx")
def export_docx(payload: ReportPayload):
    doc = Document()
    now_display = payload.now_display or datetime.now().strftime("%d/%m/%Y %I:%M %p")
    doc.add_paragraph("Dear Sir/Madam")
    doc.add_paragraph("Please check defect(s) marked against")
    doc.add_paragraph(f"Diary No:{payload.diary_no} in the matter")
    doc.add_paragraph(payload.case_title)

    main = payload.defects[: payload.max_main_rows]
    other = payload.defects[payload.max_main_rows :]

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "SlNo."
    hdr[1].text = "Defects marked during Scrutiny"
    hdr[2].text = "Date of Defects Marked"
    hdr[3].text = "Date of Defect Removed"
    for i, d in enumerate(main, start=1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = f"({d.code}) - {d.title}"
        row[2].text = d.date_marked or payload.generated_on or ""
        row[3].text = d.date_removed or ""

    if other:
        doc.add_paragraph("Any Other Defects")
        table2 = doc.add_table(rows=1, cols=2)
        hdr2 = table2.rows[0].cells
        hdr2[0].text = "Description of any other Defects:"
        hdr2[1].text = "Marked On"
        for j, d in enumerate(other, start=len(main) + 1):
            row = table2.add_row().cells
            row[0].text = f"{j}. ({d.code}) - {d.title}"
            row[1].text = d.marked_on or d.date_marked or payload.generated_on or ""

    doc.add_paragraph(f"Date :{now_display}")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="defect_report.docx"'},
    )
